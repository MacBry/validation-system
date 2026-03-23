#!/usr/bin/env python3
"""
Skrypt do wypełniania szablonu walidacji Word
"""
import sys
import json
from docx import Document

def fill_template(template_path, data_path, output_path):
    """Wypełnia szablon danymi z JSON"""
    
    # Wczytaj dane
    with open(data_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Wczytaj szablon
    doc = Document(template_path)
    
    # Cały dokument to jedna duża tabela (8 kolumn × 26 wierszy)
    if len(doc.tables) == 0:
        print("BŁĄD: Brak tabel w dokumencie!")
        return
    
    table = doc.tables[0]
    
    # ===== WYPEŁNIJ NAGŁÓWEK =====
    
    # Wiersz 1, kolumny 1-7: Dział/Pracownia
    if len(table.rows) > 1:
        for col in range(1, 8):
            table.rows[1].cells[col].text = data.get('dzial_pracownia', '')
    
    # Wiersz 3, kolumny 2-7: Nazwa urządzenia
    if len(table.rows) > 3:
        for col in range(2, 8):
            table.rows[3].cells[col].text = data.get('nazwa_urzadzenia', '')
    
    # Wiersz 4, kolumny 2-7: Data mapowania/walidacji
    if len(table.rows) > 4:
        for col in range(2, 8):
            table.rows[4].cells[col].text = data.get('data_walidacji', '')
    
    # Wiersz 5, kolumny 2-7: Numer inwentarzowy
    if len(table.rows) > 5:
        for col in range(2, 8):
            table.rows[5].cells[col].text = data.get('numer_inwentarzowy', '')
    
    # Wiersz 6, kolumny 2-7: Przechowywany materiał
    if len(table.rows) > 6:
        for col in range(2, 8):
            table.rows[6].cells[col].text = data.get('przechowywany_material', '')
    
    # ===== WYPEŁNIJ POMIARY =====
    
    pomiary = data.get('pomiary', {})
    
    # GÓRNA PÓŁKA - TYŁ LEWY
    if 'gorna_tyl_lewy' in pomiary:
        m = pomiary['gorna_tyl_lewy']
        table.rows[9].cells[1].text = m['nr_rejestratora']
        table.rows[10].cells[1].text = f"{m['temp_max']}°C"
        table.rows[11].cells[1].text = f"{m['temp_min']}°C"
        # Zaznacz X w małym kwadraciku
        table.rows[11].cells[2].text = 'X'
    
    # GÓRNA PÓŁKA - TYŁ PRAWY
    if 'gorna_tyl_prawy' in pomiary:
        m = pomiary['gorna_tyl_prawy']
        table.rows[9].cells[6].text = m['nr_rejestratora']
        table.rows[10].cells[6].text = f"{m['temp_max']}°C"
        table.rows[11].cells[6].text = f"{m['temp_min']}°C"
        # Zaznacz X w małym kwadraciku
        table.rows[11].cells[5].text = 'X'
    
    # GÓRNA PÓŁKA - PRZÓD LEWY
    if 'gorna_przod_lewy' in pomiary:
        m = pomiary['gorna_przod_lewy']
        table.rows[15].cells[1].text = m['nr_rejestratora']
        table.rows[13].cells[1].text = f"{m['temp_max']}°C"
        table.rows[14].cells[1].text = f"{m['temp_min']}°C"
        # Zaznacz X w małym kwadraciku
        table.rows[13].cells[2].text = 'X'
    
    # GÓRNA PÓŁKA - PRZÓD PRAWY
    if 'gorna_przod_prawy' in pomiary:
        m = pomiary['gorna_przod_prawy']
        table.rows[15].cells[6].text = m['nr_rejestratora']
        table.rows[13].cells[6].text = f"{m['temp_max']}°C"
        table.rows[14].cells[6].text = f"{m['temp_min']}°C"
        # Zaznacz X w małym kwadraciku
        table.rows[13].cells[5].text = 'X'
    
    # DOLNA PÓŁKA - TYŁ LEWY
    if 'dolna_tyl_lewy' in pomiary:
        m = pomiary['dolna_tyl_lewy']
        table.rows[17].cells[1].text = m['nr_rejestratora']
        table.rows[18].cells[1].text = f"{m['temp_max']}°C"
        table.rows[19].cells[1].text = f"{m['temp_min']}°C"
        # Zaznacz X w małym kwadraciku
        table.rows[19].cells[2].text = 'X'
    
    # DOLNA PÓŁKA - TYŁ PRAWY
    if 'dolna_tyl_prawy' in pomiary:
        m = pomiary['dolna_tyl_prawy']
        table.rows[17].cells[6].text = m['nr_rejestratora']
        table.rows[18].cells[6].text = f"{m['temp_max']}°C"
        table.rows[19].cells[6].text = f"{m['temp_min']}°C"
        # Zaznacz X w małym kwadraciku
        table.rows[19].cells[5].text = 'X'
    
    # DOLNA PÓŁKA - PRZÓD LEWY
    if 'dolna_przod_lewy' in pomiary:
        m = pomiary['dolna_przod_lewy']
        table.rows[23].cells[1].text = m['nr_rejestratora']
        table.rows[21].cells[1].text = f"{m['temp_max']}°C"
        table.rows[22].cells[1].text = f"{m['temp_min']}°C"
        # Zaznacz X w małym kwadraciku
        table.rows[21].cells[2].text = 'X'
    
    # DOLNA PÓŁKA - PRZÓD PRAWY
    if 'dolna_przod_prawy' in pomiary:
        m = pomiary['dolna_przod_prawy']
        table.rows[23].cells[6].text = m['nr_rejestratora']
        table.rows[21].cells[6].text = f"{m['temp_max']}°C"
        table.rows[22].cells[6].text = f"{m['temp_min']}°C"
        # Zaznacz X w małym kwadraciku
        table.rows[21].cells[5].text = 'X'
    
    # Zapisz dokument
    doc.save(output_path)
    print(f"Dokument zapisany: {output_path}")

if __name__ == '__main__':
    if len(sys.argv) != 4:
        print("Usage: python fill_validation_template.py <template.docx> <data.json> <output.docx>")
        sys.exit(1)
    
    fill_template(sys.argv[1], sys.argv[2], sys.argv[3])
